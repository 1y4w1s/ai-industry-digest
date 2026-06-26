"""
❌ Internal Server ErrorSignal - 知识库路由
提供文档上传、切片、知识图谱、对话等功能
"""

import os
import uuid
import httpx
from typing import Optional, List, Dict, Any
from datetime import datetime
from pydantic import BaseModel

from fastapi import APIRouter, File, UploadFile, Depends, HTTPException, Header, Query, BackgroundTasks

from api.models.database import get_db
from api.services.jwt_verify import verify_token, DEMO_USER_UUID
from api.services.embedding import get_embedding_service
from api.services.retrieval import get_retrieval_service
from processor.ai_processor import AIProcessor

router = APIRouter(prefix="/kb", tags=["知识库"])
db = get_db()

# 知识库对话上下文存储（独立存储，避免循环导入）
kb_chat_contexts: dict = {}

# ── 统一认证工具 ────────────────────────
# 所有 KB 路由共用同一套认证逻辑，支持：
#   - Authorization: Bearer <token>（request() 正常调用）
#   - ?token=<token>（window.open 下载等无法传 header 的场景）
#   - 临时用户兜底（未登录也能看公开文档）


async def get_current_user(
    authorization: Optional[str] = Header(None),
    token: Optional[str] = Query(None),
) -> str:
    """统一认证依赖：header 优先，query token 兜底"""
    raw = authorization or token
    if not raw:
        return DEMO_USER_UUID  # 未登录用户也能看公开文档
    
    user_id = verify_token(raw)
    if not user_id:
        return DEMO_USER_UUID  # token 无效但也允许看公开文档
    return user_id


# ── 文档配置 ────────────────────────────

SUPPORTED_EXTENSIONS = {
    ".txt": "text",
    ".md": "markdown",
    ".pdf": "pdf",
    ".docx": "docx",
}

MAX_FILE_SIZE = 10 * 1024 * 1024

# 扩展名反向映射（file_type → ext，用于安全地构建文件路径）
EXTENSION_MAP = {v: k for k, v in SUPPORTED_EXTENSIONS.items()}


def _safe_file_path(document_id: str, file_type: str) -> str:
    """安全地构建文件路径（使用服务端确定的扩展名，不信任用户输入）"""
    upload_dir = os.path.join(os.path.dirname(__file__), "..", "uploads")
    ext = EXTENSION_MAP.get(file_type, ".txt")
    return os.path.join(upload_dir, f"{document_id}{ext}")


def _doc_query(db, user_id: str):
    """构建知识库文档查询：公开文档 OR 用户自己的文档"""
    return db.client.table("kb_documents") \
        .select("*", count="exact") \
        .or_(f"is_public.eq.true,user_id.eq.{user_id}")


def _doc_access_filter(query, document_id: str, user_id: str):
    """文档访问权限过滤：公开 OR 自己拥有"""
    return query \
        .eq("id", document_id) \
        .or_(f"is_public.eq.true,user_id.eq.{user_id}")


@router.post("/documents")
async def upload_document(
    file: UploadFile = File(...),
    tags: Optional[str] = "",
    is_public: bool = True,
    user_id: str = Depends(get_current_user)
):
    """上传文档"""
    # 验证文件类型
    _, ext = os.path.splitext(file.filename.lower())
    if ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件类型，支持：{', '.join(SUPPORTED_EXTENSIONS.keys())}"
        )

    # 验证文件大小（同步读取文件内容）
    file_size = 0
    content = b""
    while True:
        chunk = file.file.read(8192)
        if not chunk:
            break
        file_size += len(chunk)
        content += chunk
        if file_size > MAX_FILE_SIZE:
            raise HTTPException(status_code=400, detail="文件大小超过限制（最大10MB）")

    # 保存文档记录
    document_id = str(uuid.uuid4())
    
    tag_list = [t.strip() for t in tags.split(",") if t.strip()] if tags else []
    
    file_type = SUPPORTED_EXTENSIONS[ext]
    safe_ext = EXTENSION_MAP[file_type]  # 服务端确定的扩展名，不信任用户输入
    
    db.client.table("kb_documents").insert({
        "id": document_id,
        "user_id": user_id,
        "name": file.filename,
        "file_type": file_type,
        "file_size": file_size,
        "status": "pending",
        "source": "user",
        "tags": tag_list,
        "is_public": is_public,
        "created_at": datetime.now().isoformat(),
        "updated_at": datetime.now().isoformat(),
    }).execute()

    # 保存文件到本地（临时）
    from api.services.document_tracker import get_document_tracker
    file_content_str = content.decode("utf-8", errors="replace")
    get_document_tracker().init_document(document_id, file_content_str)

    upload_dir = os.path.join(os.path.dirname(__file__), "..", "uploads")
    os.makedirs(upload_dir, exist_ok=True)
    file_path = os.path.join(upload_dir, f"{document_id}{safe_ext}")
    with open(file_path, "wb") as f:
        f.write(content)

    return {
        "id": document_id,
        "name": file.filename,
        "file_type": SUPPORTED_EXTENSIONS[ext],
        "file_size": file_size,
        "status": "pending",
        "tags": tag_list,
        "created_at": datetime.now().isoformat(),
    }


@router.put("/documents/{document_id}/reupload")
async def reupload_document(
    document_id: str,
    file: UploadFile = File(...),
    user_id: str = Depends(get_current_user)
):
    """重新上传文档（F-11 增量更新）
    
    检测内容是否变更：
    - 未变更 → 跳过处理，返回 "内容无变更"
    - 已变更 → 删除旧数据，全量重新处理（version+1）
    """
    from api.services.document_tracker import get_document_tracker
    
    # 验证文档属于当前用户
    doc_result = db.client.table("kb_documents") \
        .select("id, name, file_type, status") \
        .eq("id", document_id) \
        .eq("user_id", user_id) \
        .execute()
    
    if not doc_result.data:
        raise HTTPException(status_code=404, detail="文档不存在或无权限")

    # 验证文件类型
    _, ext = os.path.splitext(file.filename.lower())
    if ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件类型，支持：{', '.join(SUPPORTED_EXTENSIONS.keys())}"
        )

    # 读取文件内容
    file_size = 0
    content = b""
    while True:
        chunk = file.file.read(8192)
        if not chunk:
            break
        file_size += len(chunk)
        content += chunk
        if file_size > MAX_FILE_SIZE:
            raise HTTPException(status_code=400, detail="文件大小超过限制（最大10MB）")

    file_content_str = content.decode("utf-8", errors="replace")
    tracker = get_document_tracker()

    # 检测内容是否变更
    change = tracker.detect_change(document_id, file_content_str)
    if not change["changed"]:
        return {
            "message": change["skip_reason"],
            "document_id": document_id,
            "version": change["current_version"],
        }

    # 内容有变更：更新文档记录中的文件
    file_type = SUPPORTED_EXTENSIONS[ext]
    safe_ext = EXTENSION_MAP[file_type]
    
    upload_dir = os.path.join(os.path.dirname(__file__), "..", "uploads")
    os.makedirs(upload_dir, exist_ok=True)
    file_path = os.path.join(upload_dir, f"{document_id}{safe_ext}")
    with open(file_path, "wb") as f:
        f.write(content)

    # 更新文档记录
    db.client.table("kb_documents") \
        .update({
            "name": file.filename,
            "file_type": file_type,
            "file_size": file_size,
            "status": "processing",
            "updated_at": datetime.now().isoformat(),
        }) \
        .eq("id", document_id) \
        .execute()

    # 触发后台处理（强制重新处理）
    try:
        from tasks import process_single_document
        process_single_document.delay(document_id)
    except Exception as e:
        print(f"[KB] Celery 不可用，降级为同步处理: {e}")
        return await _sync_process_document(document_id, force_reprocess=True)

    return {
        "message": "文档更新已开始处理",
        "document_id": document_id,
        "old_version": change["current_version"],
        "new_version": change["current_version"] + 1,
    }


@router.get("/health")
async def health_check():
    """知识库服务健康检查（公开接口）"""
    return {"status": "ok", "service": "knowledge-base"}


@router.get("/documents")
async def list_documents(
    page: int = 1,
    page_size: int = 20,
    q: Optional[str] = None,
    tag: Optional[str] = None,
    status: Optional[str] = None,
    file_type: Optional[str] = None,
    source: Optional[str] = None,
    user_id: str = Depends(get_current_user)
):
    """获取文档列表"""
    
    
    query = _doc_query(db, user_id) \
        .order("created_at", desc=True)

    if q:
        query = query.ilike("name", f"%{q}%")
    if tag:
        query = query.contains("tags", [tag])
    if status:
        query = query.eq("status", status)
    if file_type:
        query = query.eq("file_type", file_type)
    if source:
        query = query.eq("source", source)

    offset = (page - 1) * page_size
    query = query.range(offset, offset + page_size - 1)

    result = query.execute()

    total = result.count or 0
    return {
        "items": result.data or [],
        "total": total,
        "page": page,
        "page_size": page_size,
        "pages": (total + page_size - 1) // page_size if page_size > 0 else 0,
    }


@router.get("/documents/{document_id}/preview")
async def preview_document(
    document_id: str,
    user_id: str = Depends(get_current_user)
):
    """预览文档内容"""
    
    
    doc_result = _doc_access_filter(
        db.client.table("kb_documents").select("id, name, file_type"),
        document_id, user_id
    ).execute()

    if not doc_result.data:
        raise HTTPException(status_code=404, detail="文档不存在")

    doc = doc_result.data[0]
    
    # 读取文件内容
    file_path = _safe_file_path(document_id, doc["file_type"])
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件不存在")
    
    content = read_file_content(file_path, doc["file_type"])
    
    return {
        "id": document_id,
        "name": doc["name"],
        "file_type": doc["file_type"],
        "content": content,
    }


@router.get("/documents/{document_id}/download")
async def download_document(
    document_id: str,
    user_id: str = Depends(get_current_user),
):
    """下载原文档（支持 ?token= 或 Authorization header 认证）"""
    from fastapi.responses import FileResponse
    
    
    
    doc_result = _doc_access_filter(
        db.client.table("kb_documents").select("id, name, file_type"),
        document_id, user_id
    ).execute()

    if not doc_result.data:
        raise HTTPException(status_code=404, detail="文档不存在")

    doc = doc_result.data[0]
    
    file_path = _safe_file_path(document_id, doc["file_type"])
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件不存在")
    
    return FileResponse(file_path, filename=doc["name"])


@router.put("/documents/{document_id}/tags")
async def update_document_tags(
    document_id: str,
    body: Dict[str, Any],
    user_id: str = Depends(get_current_user)
):
    """修改文档标签"""
    
    
    tags = body.get("tags", [])
    
    result = db.client.table("kb_documents") \
        .update({"tags": tags, "updated_at": datetime.now().isoformat()}) \
        .eq("id", document_id) \
        .eq("user_id", user_id) \
        .execute()

    if not result.data:
        raise HTTPException(status_code=404, detail="文档不存在")

    return {"message": "标签已更新", "tags": tags}


@router.post("/batch/delete")
async def batch_delete_documents(
    body: Dict[str, Any],
    user_id: str = Depends(get_current_user)
):
    """批量删除文档"""
    
    ids = body.get("ids", [])
    
    if not ids:
        raise HTTPException(status_code=400, detail="请选择要删除的文档")
    
    # 删除文档（级联删除）
    db.client.table("kb_documents") \
        .delete() \
        .in_("id", ids) \
        .eq("user_id", user_id) \
        .execute()
    
    # 删除本地文件
    upload_dir = os.path.join(os.path.dirname(__file__), "..", "uploads")
    for doc_id in ids:
        for ext in SUPPORTED_EXTENSIONS.keys():
            file_path = os.path.join(upload_dir, f"{doc_id}{ext}")
            if os.path.exists(file_path):
                os.remove(file_path)
    
    return {"message": f"已删除 {len(ids)} 个文档"}


@router.post("/batch/process")
async def batch_process_documents(
    body: Dict[str, Any],
    user_id: str = Depends(get_current_user)
):
    """批量处理文档"""
    ids = body.get("ids", [])
    
    if not ids:
        raise HTTPException(status_code=400, detail="请选择要处理的文档")
    
    results = []
    for doc_id in ids:
        try:
            result = await process_document(doc_id, user_id)
            results.append({"id": doc_id, "status": "success"})
        except Exception as e:
            results.append({"id": doc_id, "status": "failed", "error": str(e)})
    
    return {"message": "处理完成", "results": results}


@router.get("/documents/{document_id}")
async def get_document(
    document_id: str,
    user_id: str = Depends(get_current_user)
):
    """获取文档详情"""
    
    
    result = _doc_access_filter(
        db.client.table("kb_documents").select("*"),
        document_id, user_id
    ).execute()

    if not result.data:
        raise HTTPException(status_code=404, detail="文档不存在")

    return result.data[0]


@router.delete("/documents/{document_id}")
async def delete_document(
    document_id: str,
    user_id: str = Depends(get_current_user)
):
    """删除文档"""
    
    
    # 检查文档是否存在
    result = db.client.table("kb_documents") \
        .select("id") \
        .eq("id", document_id) \
        .eq("user_id", user_id) \
        .execute()

    if not result.data:
        raise HTTPException(status_code=404, detail="文档不存在")

    # 删除文档（级联删除关联的切片、实体、关系）
    db.client.table("kb_documents") \
        .delete() \
        .eq("id", document_id) \
        .eq("user_id", user_id) \
        .execute()

    # 删除本地文件
    upload_dir = os.path.join(os.path.dirname(__file__), "..", "uploads")
    for ext in SUPPORTED_EXTENSIONS.keys():
        file_path = os.path.join(upload_dir, f"{document_id}{ext}")
        if os.path.exists(file_path):
            os.remove(file_path)

    return {"message": "文档已删除"}


@router.get("/documents/{document_id}/chunks")
async def get_document_chunks(
    document_id: str,
    user_id: str = Depends(get_current_user)
):
    """获取文档切片列表"""
    
    
    # 验证文档可访问（公开 OR 自己拥有）
    doc_result = _doc_access_filter(
        db.client.table("kb_documents").select("id"),
        document_id, user_id
    ).execute()

    if not doc_result.data:
        raise HTTPException(status_code=404, detail="文档不存在")

    result = db.client.table("kb_chunks") \
        .select("*") \
        .eq("document_id", document_id) \
        .order("chunk_index") \
        .execute()

    return result.data or []


@router.get("/documents/{document_id}/graph")
async def get_document_graph(
    document_id: str,
    user_id: str = Depends(get_current_user)
):
    """获取文档知识图谱数据"""
    
    
    # 验证文档可访问（公开 OR 自己拥有）
    doc_result = _doc_access_filter(
        db.client.table("kb_documents").select("id"),
        document_id, user_id
    ).execute()

    if not doc_result.data:
        raise HTTPException(status_code=404, detail="文档不存在")

    # 获取实体
    entities = db.client.table("kb_entities") \
        .select("id, name, type") \
        .eq("document_id", document_id) \
        .execute()

    # 获取关系
    relations = db.client.table("kb_relations") \
        .select("source_entity_id, target_entity_id, relation_type, label") \
        .eq("document_id", document_id) \
        .execute()

    return {
        "nodes": entities.data or [],
        "edges": relations.data or [],
    }


@router.post("/documents/{document_id}/process")
async def process_document(
    document_id: str,
    user_id: str = Depends(get_current_user)
):
    """处理文档（切片 + Embedding + 实体识别 + 关系抽取）
    
    验证文档权限后，将实际处理交给 Celery 异步任务，
    API 立即返回，不阻塞请求。
    """
    
    # 验证文档属于用户
    doc_result = db.client.table("kb_documents") \
        .select("id, name, file_type") \
        .eq("id", document_id) \
        .eq("user_id", user_id) \
        .execute()

    if not doc_result.data:
        raise HTTPException(status_code=404, detail="文档不存在")

    # 更新状态为处理中
    db.client.table("kb_documents") \
        .update({"status": "processing", "updated_at": datetime.now().isoformat()}) \
        .eq("id", document_id) \
        .execute()

    # 触发 Celery 异步任务
    try:
        from tasks import process_single_document
        process_single_document.delay(document_id)
    except Exception as e:
        # Celery 不可用，降级为同步处理
        print(f"[KB] Celery 不可用，降级为同步处理: {e}")
        return await _sync_process_document(document_id)

    return {
        "message": "处理已开始",
        "document_id": document_id,
        "note": "文档正在后台处理，请稍后刷新页面查看结果"
    }


async def _sync_process_document(document_id: str, force_reprocess: bool = False) -> dict:
    """同步处理文档（切片 + Embedding + 实体识别 + 关系抽取）
    
    Args:
        document_id: 文档 ID
        force_reprocess: 是否强制重新处理（忽略版本检测）
    """
    from api.services.embedding import get_embedding_service
    from api.services.metadata import get_metadata_enricher
    from api.services.document_tracker import get_document_tracker
    from processor.ai_processor import AIProcessor
    
    doc_result = db.client.table("kb_documents") \
        .select("id, name, file_type, user_id") \
        .eq("id", document_id) \
        .execute()
    
    if not doc_result.data:
        raise HTTPException(status_code=404, detail="文档不存在")
    
    document = doc_result.data[0]
    file_type = document.get("file_type", "text")
    document_name = document.get("name", "")
    
    try:
        file_path = _safe_file_path(document_id, file_type)
        content = read_file_content(file_path, file_type)
        
        # F-11 增量更新：检测内容是否变更
        if not force_reprocess:
            tracker = get_document_tracker()
            change = tracker.detect_change(document_id, content)
            if not change["changed"]:
                print(f"[KB] {change['skip_reason']}")
                return {
                    "message": change["skip_reason"],
                    "document_id": document_id,
                    "chunks_count": 0,
                    "entities_count": 0,
                    "relations_count": 0,
                }
        
        # F-13 多模态支持：从 PDF/DOCX 提取图片并生成描述
        if file_type in ("pdf", "docx"):
            try:
                from api.services.image_extractor import get_image_extractor
                from api.services.image_caption import get_image_caption_service
                
                extractor = get_image_extractor()
                captioner = get_image_caption_service()
                
                if file_type == "pdf":
                    images = extractor.extract_from_pdf(file_path, document_id)
                else:
                    images = extractor.extract_from_docx(file_path, document_id)
                
                if images:
                    captioned = captioner.describe_batch(images)
                    image_lines = []
                    for img in captioned:
                        desc = img.get("description", "[图片内容]")
                        image_lines.append(f"[图片: {desc}]")
                    if image_lines:
                        content += "\n\n" + "\n\n".join(image_lines)
                        print(f"[KB] 提取了 {len(images)} 张图片，已追加到内容")
            except Exception as e:
                print(f"[KB] 图片处理失败（非致命）: {e}")
        
        chunks = split_into_chunks(content)
        
        embedding_service = get_embedding_service()
        enricher = get_metadata_enricher()
        
        # 批量生成 Embedding
        BATCH_SIZE = 10
        all_embeddings = []
        for i in range(0, len(chunks), BATCH_SIZE):
            batch = chunks[i:i + BATCH_SIZE]
            batch_embeddings = await embedding_service.get_embeddings(batch)
            all_embeddings.extend(batch_embeddings)
        
        # 保存切片（第一轮：基础元数据，不含实体）
        chunk_ids = []
        for i, (chunk, embedding) in enumerate(zip(chunks, all_embeddings)):
            chunk_id = str(uuid.uuid4())
            chunk_ids.append(chunk_id)
            
            meta = enricher.enrich(
                chunk,
                chunk_index=i,
                total_chunks=len(chunks),
                document_name=document_name,
                file_type=file_type,
                page_number=0,
            )
            
            db.client.table("kb_chunks").insert({
                "id": chunk_id,
                "document_id": document_id,
                "content": chunk,
                "chunk_index": i,
                "embedding": embedding,
                "metadata": meta,
                "created_at": datetime.now().isoformat(),
            }).execute()

        # 实体识别和关系抽取
        ai_processor = AIProcessor()
        entities, relations = await ai_processor.extract_knowledge(content)
        
        entity_map = {}
        for entity in entities:
            entity_id = str(uuid.uuid4())
            entity_map[entity["name"]] = entity_id
            db.client.table("kb_entities").insert({
                "id": entity_id,
                "document_id": document_id,
                "name": entity["name"],
                "type": entity.get("type", "concept"),
                "created_at": datetime.now().isoformat(),
            }).execute()

        for relation in relations:
            if relation["source"] in entity_map and relation["target"] in entity_map:
                db.client.table("kb_relations").insert({
                    "id": str(uuid.uuid4()),
                    "document_id": document_id,
                    "source_entity_id": entity_map[relation["source"]],
                    "target_entity_id": entity_map[relation["target"]],
                    "relation_type": relation.get("relation", "related_to"),
                    "label": relation.get("label", ""),
                    "created_at": datetime.now().isoformat(),
                }).execute()

        # 第二轮：更新切片元数据，补充实体信息
        entity_names = list(entity_map.keys())
        if entity_names:
            for chunk_id in chunk_ids:
                chunk_result = db.client.table("kb_chunks") \
                    .select("id, content, metadata") \
                    .eq("id", chunk_id) \
                    .execute()
                if chunk_result.data:
                    chunk_data = chunk_result.data[0]
                    if isinstance(chunk_data.get("metadata"), dict):
                        enriched = enricher.enrich(
                            chunk_data["content"],
                            extracted_entities=entity_names,
                        )
                        # 合并基础字段和实体字段
                        updated_meta = chunk_data["metadata"]
                        updated_meta["extracted_entities"] = enriched["extracted_entities"]
                        updated_meta["section_title"] = enriched["section_title"]
                        updated_meta["section_depth"] = enriched["section_depth"]
                        updated_meta["has_code"] = enriched["has_code"]
                        updated_meta["code_language"] = enriched["code_language"]
                        
                        db.client.table("kb_chunks") \
                            .update({"metadata": updated_meta}) \
                            .eq("id", chunk_id) \
                            .execute()

        db.client.table("kb_documents") \
            .update({
                "status": "completed",
                "chunks_count": len(chunks),
                "updated_at": datetime.now().isoformat()
            }) \
            .eq("id", document_id) \
            .execute()

        # F-11 增量更新：成功处理后递增版本号
        get_document_tracker().bump_version(document_id, content)

        return {
            "message": "处理完成",
            "chunks_count": len(chunks),
            "entities_count": len(entities),
            "relations_count": len(relations),
        }

    except Exception as e:
        db.client.table("kb_documents") \
            .update({"status": "failed", "updated_at": datetime.now().isoformat()}) \
            .eq("id", document_id) \
            .execute()
        raise HTTPException(status_code=500, detail=f"处理失败: {str(e)}")


def read_file_content(file_path: str, file_type: str, strip_html: bool = True) -> str:
    """读取文件内容（带截断保护，防止大文件超时）"""
    import time
    start = time.time()
    MAX_SECONDS = 15  # 最多处理 15 秒，超过则返回已有内容
    truncated = False

    if file_type == "text" or file_type == "markdown":
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        if strip_html:
            content = clean_html_content(content)
        return content

    elif file_type == "pdf":
        # 使用 PyMuPDF 读取 PDF（限制前 10 页）
        try:
            import fitz
            doc = fitz.open(file_path)
            total_pages = len(doc)
            max_pages = min(total_pages, 10)
            text = ""
            for i in range(max_pages):
                if time.time() - start > MAX_SECONDS:
                    truncated = True
                    break
                text += doc[i].get_text()
            doc.close()
            suffix = f"\n\n[仅显示前 {max_pages}/{total_pages} 页预览]"
            if truncated:
                suffix = f"\n\n[处理超时，仅显示前 {i+1}/{total_pages} 页]"
            return text.strip() + suffix
        except ImportError:
            raise HTTPException(status_code=500, detail="缺少 PyMuPDF 依赖")

    elif file_type == "docx":
        # 使用 python-docx 读取 DOCX（提取文本，限制 50 段，图片用占位标记）
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = []
            max_paras = 50
            for i, para in enumerate(doc.paragraphs):
                if time.time() - start > MAX_SECONDS:
                    truncated = True
                    break
                if i >= max_paras:
                    truncated = True
                    break
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            # 提取图片占位（仅统计有无，不提取图片本身）
            inline_count = 0
            for para in doc.paragraphs[:max_paras]:
                for run in para.runs:
                    if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                        inline_count += 1
            text = "\n\n".join(paragraphs)
            if inline_count > 0:
                text += f"\n\n[本文档包含 {inline_count} 张图片，文本预览中不显示图片]"
            if truncated or len(doc.paragraphs) > max_paras:
                total = len(doc.paragraphs)
                text += f"\n\n[仅显示前 {max_paras}/{total} 段文本预览]"
            return text
        except ImportError:
            raise HTTPException(status_code=500, detail="缺少 python-docx 依赖")

    else:
        raise HTTPException(status_code=400, detail=f"不支持的文件类型: {file_type}")


def clean_html_content(html: str) -> str:
    """剥离 HTML 标签，提取纯文本，修复爬取内容中的格式残留"""
    import re
    # 移除 HTML 注释
    text = re.sub(r'<!--.*?-->', '', html, flags=re.DOTALL)
    # 移除样式/脚本块
    text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)
    text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.DOTALL | re.IGNORECASE)
    # 替换 <br> <p> <div> 等块级标签为换行
    text = re.sub(r'</?(?:p|div|br|h[1-6]|blockquote|li|tr|td|th)\s*/?>', '\n', text, flags=re.IGNORECASE)
    # 移除所有剩余 HTML 标签
    text = re.sub(r'<[^>]+>', '', text)
    # HTML entity 解码
    text = text.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
    text = text.replace('&nbsp;', ' ').replace('&#x200b;', '').replace('&#8203;', '')
    text = re.sub(r'&#(\d+);', lambda m: chr(int(m.group(1))), text)
    # Reddit 特殊标记清理
    text = re.sub(r'&#x200b;', '', text)
    text = re.sub(r'/u/\w+', '', text)  # Reddit 用户名
    # 合并多余空行
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def split_into_chunks(content: str, chunk_size: int = 500, overlap: int = 80) -> List[str]:
    """将内容切分为多个切片（语义感知分割）
    
    使用 RecursiveCharacterTextSplitter 实现多层级的语义边界感知分割：
    1. 优先在段落边界（\n\n）分割 → 保证章节完整性
    2. 其次在行边界（\n）分割 → 保证行完整性
    3. 再在中文句号/问号/感叹号处分割 → 保证句子完整性
    4. 最后在字符级分割 → 兜底方案
    
    Args:
        content: 待切分的文本内容
        chunk_size: 每个切片的目标字符数（默认 500）
        overlap: 相邻切片间的重叠字符数（默认 80，约 16%）
    
    Returns:
        切分后的文本切片列表
    """
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=overlap,
            length_function=len,
            keep_separator="end",
            separators=[
                "\n\n",        # 段落边界（markdown 章节分割）
                "\n",          # 行边界
                "。",           # 中文句号
                "！",           # 中文感叹号
                "？",           # 中文问号
                "；",           # 中文分号
                "，",           # 中文逗号
                ".",            # 英文句号（含小数点和缩写）
                "!",            # 英文感叹号
                "?",            # 英文问号
                " ",            # 单词边界（英文文本）
                "",             # 字符级分割（兜底）
            ],
        )
        chunks = text_splitter.split_text(content)
        # 去除首尾空白
        return [chunk.strip() for chunk in chunks if chunk.strip()]
    except ImportError:
        # 降级方案：简单的标点边界分割
        pass
    
    # ---------- 降级实现 ----------
    chunks = []
    start = 0
    content_len = len(content)
    
    while start < content_len:
        end = min(start + chunk_size, content_len)
        if end < content_len:
            for i in range(min(overlap * 2, end - start), 0, -1):
                if content[end - i] in ".?!\n。！？":
                    end = end - i + 1
                    break
        
        chunk = content[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        next_start = end - overlap
        if next_start <= start:
            next_start = start + 1
        start = next_start
    
    return chunks


# ── 知识库对话接口 ──────────────────────────

class KBChatRequest(BaseModel):
    """知识库对话请求"""
    message: str
    document_ids: Optional[List[str]] = None  # 限制对话范围到特定文档
    session_id: Optional[str] = None  # 用于保持对话上下文


class KBChatResponse(BaseModel):
    """知识库对话响应"""
    reply: str
    session_id: str
    sources: List[Dict[str, Any]]  # 引用的知识库来源


KB_SYSTEM_PROMPT = """你是一个专业的知识库助手，帮助用户理解和查询知识库中的文档内容。

规则:
1. 你的回答必须基于提供的知识库内容，不要编造信息
2. 如果知识库中没有相关信息，请明确告知用户
3. 引用文档时使用 Markdown 链接格式：[文档名称](/knowledge?doc=文档ID)
4. 回答简洁、准确、有深度，使用中文
5. 可以跨文档整合信息，提供综合性回答"""


async def search_kb_chunks(query: str, user_id: str, document_ids: Optional[List[str]] = None, limit: int = 5) -> List[Dict[str, Any]]:
    """在知识库切片中搜索相关内容（使用高级检索）
    
    使用 AdvancedRetrievalService 实现 Query 改写 + 混合检索(向量+关键词) + RRF 融合。
    
    Args:
        query: 用户查询
        user_id: 用户ID
        document_ids: 限制搜索的文档ID列表
        limit: 返回结果数量限制
    
    Returns:
        相关切片列表，包含文档信息和切片内容
    """
    try:
        retrieval_service = get_retrieval_service()
        # 多取一些结果方便后续过滤
        results = await retrieval_service.search(query, user_id, limit * 2)
        
        # 如果指定了 document_ids，过滤
        if document_ids:
            doc_id_set = set(document_ids)
            results = [r for r in results if r.get("document", {}).get("id") in doc_id_set]
        
        return results[:limit]
    except Exception as e:
        print(f"[KB] 高级检索失败: {e}")
        return []


@router.post("/chat", response_model=KBChatResponse, tags=["知识库对话"])
async def kb_chat(
    req: KBChatRequest,
    background_tasks: BackgroundTasks,
    authorization: Optional[str] = Header(None),
    token: Optional[str] = Query(None),
):
    """知识库对话接口（带知识库上下文）
    
    上下文来源：
      1. KB_SYSTEM_PROMPT — 角色设定
      2. 相关知识库切片内容（基于关键词检索）
      3. 最近 2 轮对话历史
    
    注意：此接口独立于全局 /api/chat，专门用于知识库对话
    """
    # 认证
    raw = authorization or token
    user_id = verify_token(raw) if raw else DEMO_USER_UUID
    if not user_id:
        user_id = DEMO_USER_UUID
    
    # 检查 API Key
    api_key = os.getenv("DEEPSEEK_API_KEY")
    if not api_key:
        raise HTTPException(status_code=503, detail="AI 服务未配置（缺少 DEEPSEEK_API_KEY）")
    
    session_id = req.session_id or f"kb_session_{os.urandom(4).hex()}"
    
    # 搜索相关知识库内容
    relevant_chunks = await search_kb_chunks(
        query=req.message,
        user_id=user_id,
        document_ids=req.document_ids,
        limit=5
    )
    
    # 构建知识库上下文
    kb_context = ""
    sources = []
    
    if relevant_chunks:
        kb_context_parts = []
        for item in relevant_chunks:
            chunk = item["chunk"]
            doc = item["document"]
            
            # 添加到上下文
            kb_context_parts.append(
                f"文档：{doc.get('name', '未知文档')} (ID: {doc.get('id', '')})\n"
                f"内容：{chunk.get('content', '')}\n"
            )
            
            # 记录来源（去重）
            doc_id = doc.get('id', '')
            if doc_id and not any(s.get('id') == doc_id for s in sources):
                sources.append({
                    "id": doc_id,
                    "name": doc.get('name', ''),
                    "file_type": doc.get('file_type', ''),
                    "relevance": item["score"]
                })
        
        kb_context = "以下是知识库中的相关内容：\n\n" + "\n".join(kb_context_parts)
    else:
        kb_context = "知识库中没有找到与您问题直接相关的内容。"
    
    # 构建消息上下文
    messages = [
        {"role": "system", "content": KB_SYSTEM_PROMPT},
        {"role": "system", "content": kb_context}
    ]
    
    # 添加历史上下文（使用独立的 kb_chat_contexts）
    history = kb_chat_contexts.get(session_id, [])
    for h in history[-2:]:
        messages.append(h)
    
    # 添加当前消息
    messages.append({"role": "user", "content": req.message})
    
    # 调用 AI
    try:
        with httpx.Client(timeout=60) as client:
            resp = client.post(
                "https://api.deepseek.com/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                },
                json={
                    "model": "deepseek-chat",
                    "messages": messages,
                    "temperature": 0.3,
                    "max_tokens": 1000,
                },
            )
            resp.raise_for_status()
            data = resp.json()
            reply = data["choices"][0]["message"]["content"].strip()
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"AI 服务调用失败: {e}")
    
    # 保存上下文（使用独立的 kb_chat_contexts）
    kb_chat_contexts[session_id] = history + [
        {"role": "user", "content": req.message},
        {"role": "assistant", "content": reply},
    ]
    
    # 限制上下文大小
    if len(kb_chat_contexts) > 1000:
        keys = list(kb_chat_contexts.keys())[:500]
        for k in keys:
            del kb_chat_contexts[k]
    
    return KBChatResponse(
        reply=reply,
        session_id=session_id,
        sources=sources
    )
